# -*- coding: utf-8 -*-
"""Generate PostgreSQL API tutorial Word document for hg-microgrid-ems."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "PostgreSQL_API_教學_微電網EMS.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft JhengHei"
    font.size = Pt(11)


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text, level=0):
    return doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shading = p._element.get_or_add_pPr()
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    set_doc_defaults(doc)

    # Cover
    add_title(doc, "PostgreSQL API 入門教學")
    sub = doc.add_paragraph("連接 hg-microgrid-ems 微電網能源管理系統")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run("專案名稱：").bold = True
    info.add_run("hg-microgrid-ems\n")
    info.add_run("文件版本：").bold = True
    info.add_run("1.0\n")
    info.add_run("適用對象：").bold = True
    info.add_run("初學者")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC placeholder note
    add_heading(doc, "目錄", 1)
    for item in [
        "一、專案現況說明",
        "二、重要觀念：為何瀏覽器不能直接連 PostgreSQL",
        "三、整體架構與四步驟流程",
        "四、Step 1：安裝 PostgreSQL 並建立資料表",
        "五、Step 2：建立 API 後端（Python Flask）",
        "六、Step 3：Vue 前端呼叫 API",
        "七、Step 4：開發常見問題",
        "八、API 與頁面對照表",
        "九、建議學習順序",
        "十、名詞小字典",
        "附錄 A：完整 SQL 建表腳本",
        "附錄 B：Flask 後端完整範例",
        "附錄 C：Vue emsStore 修改範例",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    # Section 1
    add_heading(doc, "一、專案現況說明", 1)
    add_para(
        doc,
        "hg-microgrid-ems 是一套微電網能源管理系統（EMS）的前端網站，"
        "使用 Vue 3、Vite、Pinia、Element Plus 與 ECharts 建置。"
    )
    add_heading(doc, "1.1 技術棧", 2)
    add_table(
        doc,
        ["技術", "用途"],
        [
            ["Vue 3 + Vite", "網頁介面與開發工具"],
            ["Pinia（emsStore.js）", "全域狀態管理"],
            ["Element Plus", "UI 元件"],
            ["ECharts", "圖表顯示"],
        ],
    )
    add_heading(doc, "1.2 目前資料來源", 2)
    add_para(
        doc,
        "目前電力數據（太陽光電、儲能、柴發、負載、市電、SOC 等）"
        "皆由前端 emsStore.js 中的 fetchEmsData() 函式以 Math.random() 模擬產生，"
        "並非來自真實資料庫。"
    )
    add_para(doc, "主要受影響的頁面包括：", bold=True)
    for page in [
        "Dashboard.vue — 系統總覽儀表板",
        "RealTimePower.vue — 即時功率",
        "History.vue — 24 小時歷史曲線（使用 generateMockData）",
        "PVSystem.vue、ESSSystem.vue、GensetSystem.vue — 各子系統頁面",
    ]:
        add_bullet(doc, page)
    add_heading(doc, "1.3 已有的 API 測試", 2)
    add_para(
        doc,
        "專案中已有 ApiTest.vue 頁面（路由 /api-test），"
        "透過 fetch 連接後端 http://192.168.1.148:5000，"
        "測試 GET /time 與 POST /add2 兩個端點。"
        "這表示後端 API 伺服器的概念已存在，只需擴充 PostgreSQL 功能即可。"
    )

    # Section 2
    add_heading(doc, "二、重要觀念：為何瀏覽器不能直接連 PostgreSQL", 1)
    add_para(doc, "錯誤做法：", bold=True)
    add_code_block(doc, "Vue 網頁 ──直接──▶ PostgreSQL")
    add_para(doc, "正確做法：", bold=True)
    add_code_block(doc, "Vue 網頁 ──HTTP──▶ API 伺服器 ──SQL──▶ PostgreSQL")
    add_para(doc, "原因說明：", bold=True)
    for reason in [
        "瀏覽器無法執行 PostgreSQL 驅動程式（如 psycopg2、pg 模組）。",
        "資料庫帳號密碼若寫在前端 JavaScript，任何人按 F12 都能看到。",
        "需要後端負責身份驗證、權限控管與 SQL 查詢邏輯。",
    ]:
        add_bullet(doc, reason)
    add_para(doc, "正確的三層架構：", bold=True)
    add_code_block(
        doc,
        "┌─────────────────┐     fetch()      ┌──────────────────┐     SQL      ┌────────────┐\n"
        "│  Vue 前端        │ ───────────────▶ │  API 後端         │ ──────────▶ │ PostgreSQL │\n"
        "│  (port 5173)    │ ◀─────────────── │  (port 5000)      │ ◀────────── │            │\n"
        "└─────────────────┘     JSON 回應     └──────────────────┘   查詢結果    └────────────┘"
    )

    # Section 3
    add_heading(doc, "三、整體架構與四步驟流程", 1)
    add_numbered(doc, "Step 1：安裝 PostgreSQL，建立資料庫與資料表")
    add_numbered(doc, "Step 2：撰寫後端 API（Python Flask 或 Node.js Express）")
    add_numbered(doc, "Step 3：後端使用 SQL 讀寫 PostgreSQL")
    add_numbered(doc, "Step 4：Vue 前端用 fetch 呼叫 API，取代模擬數據")

    # Section 4
    add_heading(doc, "四、Step 1：安裝 PostgreSQL 並建立資料表", 1)
    add_heading(doc, "4.1 安裝 PostgreSQL", 2)
    add_para(
        doc,
        "前往 https://www.postgresql.org/download/windows/ 下載並安裝。"
        "安裝時請記住：使用者名稱（預設 postgres）、密碼、連接埠（預設 5432）。"
    )
    add_heading(doc, "4.2 建立資料庫", 2)
    add_code_block(doc, "CREATE DATABASE microgrid_ems;")
    add_heading(doc, "4.3 設計 EMS 資料表", 2)
    add_para(
        doc,
        "依 emsStore.js 的 siteOptions 與即時功率欄位，建議建立 sites（場域）"
        "與 power_readings（功率紀錄）兩張表。詳細 SQL 請見附錄 A。"
    )
    add_heading(doc, "4.4 測試查詢", 2)
    add_code_block(
        doc,
        "SELECT * FROM sites;\n"
        "SELECT * FROM power_readings ORDER BY recorded_at DESC LIMIT 10;"
    )

    # Section 5
    add_heading(doc, "五、Step 2：建立 API 後端（Python Flask）", 1)
    add_para(
        doc,
        "因 ApiTest.vue 已連接 port 5000 的 Flask 風格 API，"
        "建議在同一後端擴充 PostgreSQL 功能。"
    )
    add_heading(doc, "5.1 安裝套件", 2)
    add_code_block(doc, "pip install flask flask-cors psycopg2-binary")
    add_heading(doc, "5.2 主要 API 端點", 2)
    add_table(
        doc,
        ["方法", "路徑", "說明"],
        [
            ["GET", "/time", "回傳伺服器時間（原有測試用）"],
            ["GET", "/api/sites", "取得所有場域清單"],
            ["GET", "/api/sites/{id}/latest", "取得某場域最新即時功率"],
            ["POST", "/api/power-readings", "寫入一筆功率紀錄"],
            ["GET", "/api/sites/{id}/history?hours=24", "取得 24 小時歷史資料"],
        ],
    )
    add_para(doc, "完整 Flask 程式碼請見附錄 B。")
    add_heading(doc, "5.3 啟動後端", 2)
    add_code_block(doc, "python app.py")
    add_para(doc, "瀏覽器測試：http://localhost:5000/api/sites")

    # Section 6
    add_heading(doc, "六、Step 3：Vue 前端呼叫 API", 1)
    add_heading(doc, "6.1 環境變數設定", 2)
    add_para(doc, "在專案根目錄建立 .env.development：")
    add_code_block(doc, "VITE_API_BASE_URL=http://192.168.1.148:5000")
    add_para(doc, "在 Vue 中使用：")
    add_code_block(doc, "const API_BASE = import.meta.env.VITE_API_BASE_URL")
    add_heading(doc, "6.2 修改 emsStore.js", 2)
    add_para(
        doc,
        "將 fetchEmsData() 中的 Math.random() 模擬邏輯，"
        "改為 async fetch 呼叫 GET /api/sites/{siteId}/latest，"
        "並將回傳 JSON 寫入 pvPower、essPower 等 ref 變數。"
        "完整範例請見附錄 C。"
    )
    add_heading(doc, "6.3 修改 History.vue", 2)
    add_para(
        doc,
        "將 generateMockData() 改為呼叫 GET /api/sites/{siteId}/history?hours=24，"
        "將回傳的 rows 陣列轉換為 ECharts 所需的 timeAxis、pvData、loadData 等格式。"
    )

    # Section 7
    add_heading(doc, "七、Step 4：開發常見問題", 1)
    add_heading(doc, "7.1 CORS 跨域錯誤", 2)
    add_para(
        doc,
        "若瀏覽器 Console 出現 Access-Control-Allow-Origin 錯誤："
        "後端需安裝 flask-cors 並啟用 CORS；"
        "或在 vite.config.js 設定 proxy 將 /api 轉發至後端。"
    )
    add_code_block(
        doc,
        "// vite.config.js\n"
        "server: {\n"
        "  proxy: {\n"
        "    '/api': {\n"
        "      target: 'http://192.168.1.148:5000',\n"
        "      changeOrigin: true\n"
        "    }\n"
        "  }\n"
        "}"
    )
    add_heading(doc, "7.2 資料從哪來？", 2)
    for item in [
        "過渡期：後端定時用模擬邏輯寫入 power_readings 表",
        "正式環境：SCADA 或電表 gateway 透過 POST 寫入",
        "測試：手動 INSERT 或使用 Postman 送 POST 請求",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "7.3 安全提醒", 2)
    add_table(
        doc,
        ["項目", "建議做法"],
        [
            ["資料庫密碼", "只放在後端，使用環境變數，勿提交至 Git"],
            ["SQL 注入", "使用參數化查詢（%s），不要拼接 SQL 字串"],
            ["生產環境", "加上登入驗證、HTTPS、API 金鑰"],
        ],
    )

    # Section 8
    add_heading(doc, "八、API 與頁面對照表", 1)
    add_table(
        doc,
        ["頁面 / 功能", "目前做法", "建議 API"],
        [
            ["Dashboard 即時功率", "emsStore.fetchEmsData() 模擬", "GET /api/sites/{id}/latest"],
            ["History 24h 曲線", "generateMockData()", "GET /api/sites/{id}/history?hours=24"],
            ["場域切換", "siteOptions 寫死在 JS", "GET /api/sites"],
            ["ApiTest", "/time, /add2", "保留作連線測試"],
        ],
    )

    # Section 9
    add_heading(doc, "九、建議學習順序", 1)
    steps = [
        "使用 pgAdmin 建表，手動 INSERT 幾筆測試資料",
        "啟動 Flask 後端，用瀏覽器測試 /api/sites",
        "在 ApiTest 頁面新增卡片測試 /api/sites（熟悉 fetch 用法）",
        "修改 emsStore.fetchEmsData 接入真實 API",
        "修改 History.vue 讀取歷史資料",
    ]
    for i, s in enumerate(steps, 1):
        add_numbered(doc, s)

    # Section 10
    add_heading(doc, "十、名詞小字典", 1)
    add_table(
        doc,
        ["名詞", "說明"],
        [
            ["PostgreSQL", "開源關聯式資料庫，使用 SQL 語言存取資料"],
            ["API", "後端提供的 HTTP 網址，回傳 JSON 格式資料"],
            ["REST", "用 GET 讀取、POST 新增、PUT 更新、DELETE 刪除的 API 設計風格"],
            ["fetch", "瀏覽器內建的 HTTP 請求函式"],
            ["CORS", "跨來源資源共享，瀏覽器的跨網域安全限制機制"],
            ["Pinia store", "Vue 3 的全域狀態管理，適合集中管理 API 資料"],
            ["JSON", "JavaScript Object Notation，前後端交換資料的標準格式"],
        ],
    )

    doc.add_page_break()

    # Appendix A
    add_heading(doc, "附錄 A：完整 SQL 建表腳本", 1)
    add_code_block(
        doc,
        "-- 場域基本資料\n"
        "CREATE TABLE sites (\n"
        "  id              VARCHAR(50) PRIMARY KEY,\n"
        "  name            VARCHAR(100) NOT NULL,\n"
        "  location        VARCHAR(100),\n"
        "  pv_capacity_kw  NUMERIC(10,2),\n"
        "  ess_power_kw    NUMERIC(10,2),\n"
        "  ess_energy_kwh  NUMERIC(10,2),\n"
        "  gen_capacity_kw NUMERIC(10,2),\n"
        "  load_base_kw    NUMERIC(10,2)\n"
        ");\n\n"
        "-- 即時／歷史功率紀錄\n"
        "CREATE TABLE power_readings (\n"
        "  id          SERIAL PRIMARY KEY,\n"
        "  site_id     VARCHAR(50) REFERENCES sites(id),\n"
        "  recorded_at TIMESTAMPTZ DEFAULT NOW(),\n"
        "  pv_power    NUMERIC(10,2),\n"
        "  ess_power   NUMERIC(10,2),\n"
        "  gen_power   NUMERIC(10,2),\n"
        "  load_power  NUMERIC(10,2),\n"
        "  grid_power  NUMERIC(10,2),\n"
        "  soc         NUMERIC(5,2),\n"
        "  mode        SMALLINT\n"
        ");\n\n"
        "INSERT INTO sites VALUES\n"
        "  ('jiali-junior-high', '臺南市佳里國中後港校區', '臺南市佳里區',\n"
        "   132, 600, 1316, 200, 70),\n"
        "  ('ruifeng-elementary', '臺南市瑞峰國小', '臺南市南化區',\n"
        "   91.12, 600, 1200, 60, 30),\n"
        "  ('zengwen-vision-park', '臺南市曾文市政願景園區', '臺南市麻豆區',\n"
        "   497.39, 1000, 1800, 200, 76);"
    )

    # Appendix B
    add_heading(doc, "附錄 B：Flask 後端完整範例", 1)
    add_code_block(
        doc,
        "from flask import Flask, jsonify, request\n"
        "from flask_cors import CORS\n"
        "import psycopg2\n"
        "from psycopg2.extras import RealDictCursor\n"
        "from datetime import datetime\n\n"
        "app = Flask(__name__)\n"
        "CORS(app)\n\n"
        "DB_CONFIG = {\n"
        '    "host": "localhost",\n'
        '    "port": 5432,\n'
        '    "database": "microgrid_ems",\n'
        '    "user": "postgres",\n'
        '    "password": "你的密碼"\n'
        "}\n\n"
        "def get_db():\n"
        "    return psycopg2.connect(**DB_CONFIG, cursor_factory=RealDictCursor)\n\n"
        '@app.route("/time")\n'
        "def get_time():\n"
        '    return jsonify({"time": datetime.now().isoformat()})\n\n'
        '@app.route("/api/sites")\n'
        "def list_sites():\n"
        "    conn = get_db()\n"
        "    cur = conn.cursor()\n"
        '    cur.execute("SELECT * FROM sites ORDER BY name")\n'
        "    rows = cur.fetchall()\n"
        "    cur.close(); conn.close()\n"
        "    return jsonify(rows)\n\n"
        '@app.route("/api/sites/<site_id>/latest")\n'
        "def latest_power(site_id):\n"
        "    conn = get_db()\n"
        "    cur = conn.cursor()\n"
        "    cur.execute(\"\"\"\n"
        "        SELECT * FROM power_readings\n"
        "        WHERE site_id = %s ORDER BY recorded_at DESC LIMIT 1\n"
        '    """, (site_id,))\n'
        "    row = cur.fetchone()\n"
        "    cur.close(); conn.close()\n"
        "    if not row:\n"
        '        return jsonify({"error": "尚無資料"}), 404\n'
        "    return jsonify(row)\n\n"
        '@app.route("/api/power-readings", methods=["POST"])\n'
        "def create_reading():\n"
        "    data = request.get_json()\n"
        "    conn = get_db()\n"
        "    cur = conn.cursor()\n"
        "    cur.execute(\"\"\"\n"
        "        INSERT INTO power_readings\n"
        "          (site_id, pv_power, ess_power, gen_power,\n"
        "           load_power, grid_power, soc, mode)\n"
        "        VALUES (%s,%s,%s,%s,%s,%s,%s,%s) RETURNING *\n"
        "    \"\"\", (data['site_id'], data['pv_power'], data['ess_power'],\n"
        "            data['gen_power'], data['load_power'], data['grid_power'],\n"
        "            data['soc'], data['mode']))\n"
        "    row = cur.fetchone()\n"
        "    conn.commit()\n"
        "    cur.close(); conn.close()\n"
        "    return jsonify(row), 201\n\n"
        '@app.route("/api/sites/<site_id>/history")\n'
        "def history_power(site_id):\n"
        '    hours = request.args.get("hours", 24, type=int)\n'
        "    conn = get_db()\n"
        "    cur = conn.cursor()\n"
        "    cur.execute(\"\"\"\n"
        "        SELECT recorded_at, pv_power, ess_power, gen_power,\n"
        "               load_power, grid_power, soc, mode\n"
        "        FROM power_readings\n"
        "        WHERE site_id = %s\n"
        "          AND recorded_at >= NOW() - INTERVAL '%s hours'\n"
        "        ORDER BY recorded_at ASC\n"
        '    """, (site_id, hours))\n'
        "    rows = cur.fetchall()\n"
        "    cur.close(); conn.close()\n"
        "    return jsonify(rows)\n\n"
        'if __name__ == "__main__":\n'
        '    app.run(host="0.0.0.0", port=5000, debug=True)'
    )

    # Appendix C
    add_heading(doc, "附錄 C：Vue emsStore 修改範例", 1)
    add_code_block(
        doc,
        "const API_BASE = import.meta.env.VITE_API_BASE_URL || 'http://localhost:5000'\n\n"
        "const fetchEmsData = async () => {\n"
        "  try {\n"
        "    const siteId = selectedSiteId.value\n"
        "    const res = await fetch(`${API_BASE}/api/sites/${siteId}/latest`)\n"
        "    if (!res.ok) throw new Error(`HTTP ${res.status}`)\n"
        "    const data = await res.json()\n\n"
        "    pvPower.value    = Number(data.pv_power)\n"
        "    essPower.value   = Number(data.ess_power)\n"
        "    genPower.value   = Number(data.gen_power)\n"
        "    loadPower.value  = Number(data.load_power)\n"
        "    gridPower.value  = Number(data.grid_power)\n"
        "    soc.value        = Number(data.soc)\n"
        "    currentMode.value = Number(data.mode)\n"
        "  } catch (err) {\n"
        "    console.error('無法取得 EMS 數據', err)\n"
        "  }\n"
        "}"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"已產生：{OUTPUT}")


if __name__ == "__main__":
    build_document()
